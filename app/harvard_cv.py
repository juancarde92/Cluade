"""Rewrite a resume into Harvard-style content and render it as a .docx."""
import copy
import json
import os
import re

from docx import Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

ASSETS_DIR = os.path.join(os.path.dirname(__file__), "assets", "harvard_templates")

TEMPLATE_CHOICES = {
    "moderno": {
        "label": "Moderno",
        "description": "Diseño limpio de una columna, encabezados con línea inferior.",
        "path": os.path.join(ASSETS_DIR, "moderno.docx"),
    },
    "oficial": {
        "label": "Harvard Oficial",
        "description": "Plantilla clásica del Harvard Office of Career Services.",
        "path": os.path.join(ASSETS_DIR, "oficial.docx"),
    },
}

# Fixed section-header translations for the two bundled templates. Kept as a
# static table (not AI-generated) so headers stay consistent and don't
# depend on the model echoing an exact schema back.
LANGUAGES = {
    "es": {
        "label": "Español",
        "prompt_name": "español",
        "moderno_headers": {
            "experience": "EXPERIENCIA PROFESIONAL", "education": "EDUCACIÓN",
            "skills": "SKILLS ADICIONALES", "technologies": "TECNOLOGÍAS",
        },
        "oficial_headers": {"education": "Educación", "experience": "Experiencia",
                             "skills_heading": "Habilidades e Intereses", "technical": "Técnico: "},
        "linkedin": {"title": "Tu Perfil de LinkedIn Optimizado", "headline": "Titular",
                     "about": "Acerca de", "recommendations": "Recomendaciones para tu perfil"},
    },
    "en": {
        "label": "English",
        "prompt_name": "English",
        "moderno_headers": {
            "experience": "PROFESSIONAL EXPERIENCE", "education": "EDUCATION",
            "skills": "ADDITIONAL SKILLS", "technologies": "TECHNOLOGIES",
        },
        "oficial_headers": {"education": "Education", "experience": "Experience",
                             "skills_heading": "Skills & Interests", "technical": "Technical: "},
        "linkedin": {"title": "Your Optimized LinkedIn Profile", "headline": "Headline",
                     "about": "About", "recommendations": "Recommendations for your profile"},
    },
    "fr": {
        "label": "Français",
        "prompt_name": "français",
        "moderno_headers": {
            "experience": "EXPÉRIENCE PROFESSIONNELLE", "education": "FORMATION",
            "skills": "COMPÉTENCES SUPPLÉMENTAIRES", "technologies": "TECHNOLOGIES",
        },
        "oficial_headers": {"education": "Formation", "experience": "Expérience",
                             "skills_heading": "Compétences et intérêts", "technical": "Techniques : "},
        "linkedin": {"title": "Votre profil LinkedIn optimisé", "headline": "Titre",
                     "about": "À propos", "recommendations": "Recommandations pour votre profil"},
    },
    "pt": {
        "label": "Português",
        "prompt_name": "português",
        "moderno_headers": {
            "experience": "EXPERIÊNCIA PROFISSIONAL", "education": "FORMAÇÃO ACADÊMICA",
            "skills": "HABILIDADES ADICIONAIS", "technologies": "TECNOLOGIAS",
        },
        "oficial_headers": {"education": "Formação", "experience": "Experiência",
                             "skills_heading": "Habilidades e Interesses", "technical": "Técnico: "},
        "linkedin": {"title": "Seu Perfil do LinkedIn Otimizado", "headline": "Título",
                     "about": "Sobre", "recommendations": "Recomendações para seu perfil"},
    },
}
DEFAULT_LANGUAGE = "es"

SYSTEM_PROMPT = (
    "Eres un experto en redacción de hojas de vida en el formato estándar de Harvard "
    "(Harvard Office of Career Services). Reescribes el contenido de un CV existente: "
    "mejoras la redacción, usas verbos de acción fuertes al inicio de cada logro, "
    "cuantificas resultados cuando sea razonable inferirlo del texto original (sin "
    "inventar cifras que no estén sugeridas), y mantienes un tono profesional y conciso. "
    "No inventes experiencia, empresas, cargos o títulos que no existan en el CV original; "
    "solo mejora la redacción y organización de lo que ya está ahí."
)

RESPONSE_SCHEMA_HINT = """
Responde ÚNICAMENTE con un JSON válido (sin markdown, sin explicación adicional) con esta forma exacta:
{
  "name": "Nombre completo",
  "contact_line": "correo · teléfono · ciudad · linkedin",
  "summary": "2-3 líneas de resumen profesional",
  "education": [
    {"school": "Universidad", "location": "Ciudad, País", "degree": "Título obtenido", "dates": "Año inicio - Año fin"}
  ],
  "experience": [
    {"organization": "Empresa", "location": "Ciudad, País", "title": "Cargo", "dates": "Mes Año - Mes Año",
     "bullets": ["Logro reescrito 1", "Logro reescrito 2"]}
  ],
  "skills": "Lista de habilidades y herramientas separadas por comas"
}
"""


def _call_gemini(system_prompt, user_prompt):
    # Calls the Gemini REST API directly with `requests` instead of the
    # google-generativeai SDK. The SDK pulls in grpcio/protobuf/google-api-core,
    # a heavy dependency tree that was pushing memory past low-RAM deployment
    # limits (e.g. Render's 512MB free tier) even with transport="rest".
    import requests

    api_key = os.environ["GEMINI_API_KEY"]
    model_name = os.environ.get("GEMINI_MODEL", "gemini-2.5-flash")
    url = f"https://generativelanguage.googleapis.com/v1beta/models/{model_name}:generateContent"

    payload = {
        "system_instruction": {"parts": [{"text": system_prompt}]},
        "contents": [{"parts": [{"text": user_prompt}]}],
    }
    resp = requests.post(url, params={"key": api_key}, json=payload, timeout=90)
    resp.raise_for_status()
    body = resp.json()
    text = body["candidates"][0]["content"]["parts"][0]["text"].strip()
    text = re.sub(r"^```(?:json)?|```$", "", text.strip(), flags=re.MULTILINE).strip()
    return json.loads(text)


def rewrite_resume_harvard(raw_text, language_code=DEFAULT_LANGUAGE):
    lang = LANGUAGES.get(language_code, LANGUAGES[DEFAULT_LANGUAGE])
    prompt = (
        f"Aquí está el texto extraído de un CV (puede tener errores de formato "
        f"por la extracción):\n\n---\n{raw_text}\n---\n\n"
        f"IMPORTANTE: Escribe TODO el contenido del JSON de salida (resumen, títulos de cargos, "
        f"logros, títulos de estudios, habilidades) en {lang['prompt_name']}, sin importar en qué "
        f"idioma esté el CV original. No traduzcas nombres propios de personas, empresas ni "
        f"instituciones educativas.\n\n{RESPONSE_SCHEMA_HINT}"
    )
    return _call_gemini(SYSTEM_PROMPT, prompt)


LINKEDIN_SYSTEM_PROMPT = (
    "Eres un experto en marca personal y optimización de perfiles de LinkedIn. A partir de un CV, "
    "redactas un titular y una sección 'Acerca de' atractivos y honestos (sin inventar logros que "
    "no estén respaldados por el CV), y das recomendaciones prácticas y accionables para mejorar "
    "la visibilidad y el impacto del perfil (no son elogios ni testimonios de terceros, son consejos "
    "de optimización)."
)

LINKEDIN_SCHEMA_HINT = """
Responde ÚNICAMENTE con un JSON válido (sin markdown) con esta forma exacta:
{
  "headline": "Titular profesional para LinkedIn, máximo 220 caracteres",
  "about": "Sección 'Acerca de' en primera persona, 3-4 párrafos cortos",
  "recommendations": ["Recomendación práctica 1", "Recomendación práctica 2", "..."]
}
Incluye entre 5 y 7 recomendaciones sobre: foto de perfil, banner, palabras clave del titular,
frecuencia de publicación, participación en la red, y cómo pedir recomendaciones genuinas a
antiguos colegas o jefes.
"""


def generate_linkedin_profile(raw_text, language_code=DEFAULT_LANGUAGE):
    lang = LANGUAGES.get(language_code, LANGUAGES[DEFAULT_LANGUAGE])
    prompt = (
        f"Aquí está el texto extraído de un CV:\n\n---\n{raw_text}\n---\n\n"
        f"IMPORTANTE: Escribe todo el contenido en {lang['prompt_name']}.\n\n{LINKEDIN_SCHEMA_HINT}"
    )
    return _call_gemini(LINKEDIN_SYSTEM_PROMPT, prompt)


def _as_list(value):
    if isinstance(value, list):
        return [v for v in value if isinstance(v, dict)]
    if isinstance(value, dict):
        return [value]
    return []


# ---------------------------------------------------------------------------
# Generic docx XML helpers used to fill the two uploaded Harvard templates
# while preserving their original design (fonts, spacing, bullets, tab
# stops for right-aligned dates, table layout).
# ---------------------------------------------------------------------------

def _distinct_cells(row):
    """A merged cell is returned once per spanned column by python-docx;
    dedupe by identity to get the real, independent cells in a row."""
    seen_ids = set()
    cells = []
    for cell in row.cells:
        if id(cell) not in seen_ids:
            seen_ids.add(id(cell))
            cells.append(cell)
    return cells


def _set_para_text(paragraph, text):
    """Overwrite a paragraph's visible text with a single new run, preserving
    formatting from the first existing run found. Removes both plain runs
    (w:r) and hyperlink-wrapped runs (w:hyperlink > w:r) — paragraph.runs
    only sees the former, so a leftover hyperlink's text would otherwise
    stay stuck onto the new text."""
    p = paragraph._p
    rpr_template = None
    for child in list(p):
        if child.tag == qn("w:r"):
            if rpr_template is None:
                rpr = child.find(qn("w:rPr"))
                if rpr is not None:
                    rpr_template = copy.deepcopy(rpr)
            p.remove(child)
        elif child.tag == qn("w:hyperlink"):
            if rpr_template is None:
                inner_r = child.find(qn("w:r"))
                rpr = inner_r.find(qn("w:rPr")) if inner_r is not None else None
                if rpr is not None:
                    rpr_template = copy.deepcopy(rpr)
            p.remove(child)

    run = paragraph.add_run(text)
    if rpr_template is not None:
        existing = run._r.find(qn("w:rPr"))
        if existing is not None:
            run._r.remove(existing)
        run._r.insert(0, rpr_template)


def _duplicate_paragraph(ref_paragraph):
    new_p = copy.deepcopy(ref_paragraph._p)
    ref_paragraph._p.addnext(new_p)
    return Paragraph(new_p, ref_paragraph._parent)


def _remove_paragraph(paragraph):
    el = paragraph._p
    el.getparent().remove(el)


def _find_paragraph(doc, text):
    for p in doc.paragraphs:
        if p.text.strip() == text:
            return p
    return None


def _set_bullet_list(cell_or_doc, items):
    """Resize a run of paragraphs (inside a table cell, or in the document
    body) to match `items`, reusing the first paragraph as the formatting
    template (preserves native Word bullet numbering)."""
    items = [str(i).strip() for i in (items or []) if str(i or "").strip()]
    paragraphs = cell_or_doc.paragraphs
    if not paragraphs:
        return
    template_p = paragraphs[0]
    for extra in paragraphs[1:]:
        _remove_paragraph(extra)
    if not items:
        _set_para_text(template_p, "")
        return
    _set_para_text(template_p, items[0])
    last_p = template_p
    for item in items[1:]:
        last_p = _duplicate_paragraph(last_p)
        _set_para_text(last_p, item)


def _run_rpr_template(paragraph, bold):
    for run in paragraph.runs:
        if bool(run.bold) == bold and run.text.strip():
            rpr = run._r.find(qn("w:rPr"))
            return copy.deepcopy(rpr) if rpr is not None else None
    return None


def _set_tab_line(paragraph, left_text, right_text):
    """Replace a paragraph's content with `left \\t right`, reusing the
    paragraph's existing bold/normal run formatting (used for the
    Organization/City and Position/Dates lines in the official template)."""
    bold_rpr = _run_rpr_template(paragraph, bold=True)
    normal_rpr = _run_rpr_template(paragraph, bold=False)
    for run in list(paragraph.runs):
        run._r.getparent().remove(run._r)

    def add(text, rpr):
        run = paragraph.add_run(text)
        if rpr is not None:
            existing = run._r.find(qn("w:rPr"))
            if existing is not None:
                run._r.remove(existing)
            run._r.insert(0, copy.deepcopy(rpr))
        return run

    add(left_text, bold_rpr)
    add("\t", normal_rpr)
    add(right_text, normal_rpr)


def _remove_row(row):
    row._tr.getparent().remove(row._tr)


def _find_header_row(table, text):
    for row in table.rows:
        cells = _distinct_cells(row)
        if cells and cells[0].paragraphs and cells[0].paragraphs[0].text.strip() == text:
            return row
    raise ValueError(f"header row not found: {text!r}")


def _rows_between(table, start_row, end_row):
    result = []
    el = start_row._tr.getnext()
    while el is not None and el is not end_row._tr:
        for r in table.rows:
            if r._tr is el:
                result.append(r)
                break
        el = el.getnext()
    return result


def _rows_after(table, start_row):
    result = []
    el = start_row._tr.getnext()
    while el is not None:
        for r in table.rows:
            if r._tr is el:
                result.append(r)
                break
        el = el.getnext()
    return result


def build_from_moderno(data, output_path, language_code=DEFAULT_LANGUAGE):
    if not isinstance(data, dict):
        data = {}
    headers = LANGUAGES.get(language_code, LANGUAGES[DEFAULT_LANGUAGE])["moderno_headers"]
    doc = Document(TEMPLATE_CHOICES["moderno"]["path"])
    table = doc.tables[0]

    name_cell = _distinct_cells(table.rows[0])[0]
    _set_para_text(name_cell.paragraphs[0], str(data.get("name") or "").strip())

    contact_cell = _distinct_cells(table.rows[1])[0]
    _set_para_text(contact_cell.paragraphs[0], str(data.get("contact_line") or "").strip())

    summary_cell = _distinct_cells(table.rows[3])[0]
    _set_para_text(summary_cell.paragraphs[0], str(data.get("summary") or "").strip())

    # Locate each header row by the text baked into the bundled template
    # (always Spanish), then relabel it to the requested output language.
    exp_header = _find_header_row(table, "EXPERIENCIA PROFESIONAL")
    edu_header = _find_header_row(table, "EDUCACIÓN")
    skills_header = _find_header_row(table, "SKILLS ADICIONALES")
    tech_header = _find_header_row(table, "TECNOLOGÍAS")
    for row, key in ((exp_header, "experience"), (edu_header, "education"),
                     (skills_header, "skills"), (tech_header, "technologies")):
        _set_para_text(_distinct_cells(row)[0].paragraphs[0], headers[key])

    # --- Experience: groups of 3 rows (title/loc, bullets, spacer) ---
    between = _rows_between(table, exp_header, edu_header)
    groups = [between[i:i + 3] for i in range(1, len(between), 3)]
    groups = [g for g in groups if len(g) == 3]

    experience = _as_list(data.get("experience"))
    target = max(len(experience), 1)
    while len(groups) < target:
        title_row, bullets_row, spacer_row = groups[-1]
        new_title = _duplicate_row_after(table, title_row, spacer_row)
        new_bullets = _duplicate_row_after(table, bullets_row, new_title)
        new_spacer = _duplicate_row_after(table, spacer_row, new_bullets)
        groups.append([new_title, new_bullets, new_spacer])
    while len(groups) > target:
        for row in groups.pop():
            _remove_row(row)

    filled_experience = experience if experience else [{}]
    for (title_row, bullets_row, _spacer), job in zip(groups, filled_experience):
        cells = _distinct_cells(title_row)
        cell0, cell1 = cells[0], cells[1]
        _set_para_text(cell0.paragraphs[0], str(job.get("organization") or ""))
        if len(cell0.paragraphs) > 1:
            _set_para_text(cell0.paragraphs[1], str(job.get("title") or ""))
        _set_para_text(cell1.paragraphs[0], str(job.get("location") or ""))
        if len(cell1.paragraphs) > 1:
            _set_para_text(cell1.paragraphs[1], str(job.get("dates") or ""))

        bullets = job.get("bullets") or []
        if not isinstance(bullets, list):
            bullets = [bullets]
        _set_bullet_list(_distinct_cells(bullets_row)[0], bullets)

    # --- Education: groups of 2 rows (entry, spacer) ---
    edu_between = _rows_between(table, edu_header, skills_header)
    edu_groups = [edu_between[i:i + 2] for i in range(1, len(edu_between), 2)]
    edu_groups = [g for g in edu_groups if len(g) == 2]

    education = _as_list(data.get("education"))
    target = max(len(education), 1)
    while len(edu_groups) < target:
        entry_row, spacer_row = edu_groups[-1]
        new_entry = _duplicate_row_after(table, entry_row, spacer_row)
        new_spacer = _duplicate_row_after(table, spacer_row, new_entry)
        edu_groups.append([new_entry, new_spacer])
    while len(edu_groups) > target:
        for row in edu_groups.pop():
            _remove_row(row)

    filled_education = education if education else [{}]
    for (entry_row, _spacer), edu in zip(edu_groups, filled_education):
        cells = _distinct_cells(entry_row)
        cell0, cell1 = cells[0], cells[1]
        paras0 = cell0.paragraphs
        _set_para_text(paras0[0], str(edu.get("school") or ""))
        if len(paras0) > 1:
            _set_para_text(paras0[1], str(edu.get("degree") or ""))
        for extra in paras0[2:]:
            _remove_paragraph(extra)
        paras1 = cell1.paragraphs
        _set_para_text(paras1[0], str(edu.get("location") or ""))
        if len(paras1) > 1:
            _set_para_text(paras1[1], str(edu.get("dates") or ""))

    # --- Skills / Technologies ---
    skills_raw = data.get("skills")
    if isinstance(skills_raw, list):
        skills_items = [str(s).strip() for s in skills_raw if str(s).strip()]
    else:
        skills_items = [s.strip() for s in str(skills_raw or "").split(",") if s.strip()]

    skills_between = _rows_between(table, skills_header, tech_header)
    if skills_between:
        _set_bullet_list(_distinct_cells(skills_between[-1])[0], skills_items)

    tech_rows = _rows_after(table, tech_header)
    if tech_rows:
        tech_cell = _distinct_cells(tech_rows[-1])[0]
        _set_para_text(tech_cell.paragraphs[0], ", ".join(skills_items))
        for extra in tech_cell.paragraphs[1:]:
            _remove_paragraph(extra)

    doc.save(output_path)


def _duplicate_row_after(table, ref_row, after_row):
    new_tr = copy.deepcopy(ref_row._tr)
    after_row._tr.addnext(new_tr)
    for r in table.rows:
        if r._tr is new_tr:
            return r
    raise RuntimeError("could not locate duplicated row")


def _paragraphs_between(doc, start_text, end_text):
    started = False
    result = []
    for p in doc.paragraphs:
        t = p.text.strip()
        if not started:
            if t == start_text:
                started = True
            continue
        if t == end_text:
            break
        result.append(p)
    return result


def _duplicate_paragraph_after(ref_paragraph, after_paragraph):
    new_p = copy.deepcopy(ref_paragraph._p)
    after_paragraph._p.addnext(new_p)
    return Paragraph(new_p, ref_paragraph._parent)


def _set_labeled_line(paragraph, label, value):
    bold_rpr = _run_rpr_template(paragraph, bold=True)
    for run in list(paragraph.runs):
        run._r.getparent().remove(run._r)

    def add(text, rpr):
        run = paragraph.add_run(text)
        if rpr is not None:
            existing = run._r.find(qn("w:rPr"))
            if existing is not None:
                run._r.remove(existing)
            run._r.insert(0, copy.deepcopy(rpr))

    add(label, bold_rpr)
    add(value, None)


def build_from_oficial(data, output_path, language_code=DEFAULT_LANGUAGE):
    if not isinstance(data, dict):
        data = {}
    headers = LANGUAGES.get(language_code, LANGUAGES[DEFAULT_LANGUAGE])["oficial_headers"]
    doc = Document(TEMPLATE_CHOICES["oficial"]["path"])

    name_p = _find_paragraph(doc, "Firstname Lastname")
    if name_p is not None:
        _set_para_text(name_p, str(data.get("name") or "").strip())
        if name_p.runs:
            name_p.runs[0].bold = True

    contact_p = _find_paragraph(
        doc, "Home or Campus Street Address • City, State Zip • youremail@college.harvard.edu • phone number"
    )
    if contact_p is not None:
        _set_para_text(contact_p, str(data.get("contact_line") or "").strip())

    # Keep references to the section headings now (matched by the template's
    # original English text) so they can be relabeled at the end, after
    # they've been used as text-based section boundaries below.
    education_heading_p = _find_paragraph(doc, "Education")
    experience_heading_p = _find_paragraph(doc, "Experience")

    # --- Education: drop the coursework/study-abroad/high-school placeholders ---
    for text in (
        "Relevant Coursework: [Note: Optional. Awards and honors can also be listed here.]",
        "Study Abroad [Note: If Applicable]\tCity, Country",
        "Study abroad coursework in \t.\tMonth Year – Month Year",
        "High School Name\tCity, State",
        "[Note: May include GPA, SAT/ACT scores, or academic honors an employer may want to know]\t Graduation Date",
    ):
        p = _find_paragraph(doc, text)
        if p is not None:
            _remove_paragraph(p)

    edu_school_p = _find_paragraph(doc, "Harvard University\tCambridge, MA")
    edu_degree_p = _find_paragraph(
        doc, "Degree, Concentration. GPA [Note: GPA is Optional] \tGraduation Date \nThesis [Note: Optional]"
    )
    education = _as_list(data.get("education"))
    if edu_school_p is not None and edu_degree_p is not None:
        if not education:
            _remove_paragraph(edu_degree_p)
            _remove_paragraph(edu_school_p)
        else:
            school_p, degree_p = edu_school_p, edu_degree_p
            for edu in education:
                _set_tab_line(school_p, str(edu.get("school") or ""), str(edu.get("location") or ""))
                _set_tab_line(degree_p, str(edu.get("degree") or ""), str(edu.get("dates") or ""))
                if edu is not education[-1]:
                    new_school = _duplicate_paragraph_after(school_p, degree_p)
                    new_degree = _duplicate_paragraph_after(degree_p, new_school)
                    school_p, degree_p = new_school, new_degree

    # --- Experience: groups of (org line, position line, N bullet lines).
    # The template has blank spacer paragraphs between groups, so group
    # boundaries are detected by paragraph style rather than a fixed size:
    # an org/position line is Normal style, bullets are "List Paragraph".
    exp_section = _paragraphs_between(doc, "Experience", "Leadership & Activities")
    job_groups = []
    i, n = 0, len(exp_section)
    while i < n:
        p = exp_section[i]
        if not p.text.strip():
            i += 1
            continue
        if i + 1 >= n:
            break
        org_p, pos_p = p, exp_section[i + 1]
        i += 2
        bullets = []
        while i < n and (exp_section[i].style.name if exp_section[i].style else "") == "List Paragraph":
            bullets.append(exp_section[i])
            i += 1
        job_groups.append([org_p, pos_p] + bullets)

    experience = _as_list(data.get("experience"))
    target = max(len(experience), 1)
    while len(job_groups) < target:
        last_group = job_groups[-1]
        new_group = []
        insert_after = last_group[-1]
        for ref_p in last_group:
            new_p = _duplicate_paragraph_after(ref_p, insert_after)
            new_group.append(new_p)
            insert_after = new_p
        job_groups.append(new_group)
    while len(job_groups) > target:
        for p in job_groups.pop():
            _remove_paragraph(p)

    filled_experience = experience if experience else [{}]
    for group, job in zip(job_groups, filled_experience):
        org_p, pos_p, *bullet_paras = group
        _set_tab_line(org_p, str(job.get("organization") or ""), str(job.get("location") or ""))
        _set_tab_line(pos_p, str(job.get("title") or ""), str(job.get("dates") or ""))

        bullets = job.get("bullets") or []
        if not isinstance(bullets, list):
            bullets = [bullets]
        while bullets and len(bullet_paras) < len(bullets):
            bullet_paras.append(_duplicate_paragraph(bullet_paras[-1]))
        while len(bullet_paras) > max(len(bullets), 0):
            _remove_paragraph(bullet_paras.pop())
        for bp, text in zip(bullet_paras, bullets):
            _set_para_text(bp, text)

    # --- Leadership & Activities: not part of our data model, drop it ---
    leadership_heading = _find_paragraph(doc, "Leadership & Activities")
    leadership_section = _paragraphs_between(doc, "Leadership & Activities", "Skills & Interests [Note: Optional]")
    for p in leadership_section:
        _remove_paragraph(p)
    if leadership_heading is not None:
        _remove_paragraph(leadership_heading)

    # --- Skills & Interests ---
    skills_heading = _find_paragraph(doc, "Skills & Interests [Note: Optional]")
    if skills_heading is not None:
        for run in list(skills_heading.runs):
            run._r.getparent().remove(run._r)
        run = skills_heading.add_run(headers["skills_heading"])
        run.bold = True

    skills_raw = data.get("skills")
    skills_text = ", ".join(str(s) for s in skills_raw) if isinstance(skills_raw, list) else str(skills_raw or "")

    technical_p = _find_paragraph(
        doc, "Technical: List computer software and programming languages and your level of fluency"
    )
    if technical_p is not None:
        _set_labeled_line(technical_p, headers["technical"], skills_text)

    for text in (
        "Language: List foreign languages and your level of fluency",
        "Laboratory: List scientific / research lab techniques or tools [If Applicable]",
        "Interests: List activities you enjoy that may spark interview conversation",
    ):
        p = _find_paragraph(doc, text)
        if p is not None:
            _remove_paragraph(p)

    # Relabel Education/Experience headings now that they're no longer
    # needed as English text anchors elsewhere in this function.
    if education_heading_p is not None:
        _set_para_text(education_heading_p, headers["education"])
    if experience_heading_p is not None:
        _set_para_text(experience_heading_p, headers["experience"])

    doc.save(output_path)


TEMPLATE_BUILDERS = {
    "moderno": build_from_moderno,
    "oficial": build_from_oficial,
}


def build_docx(template_key, data, output_path, language_code=DEFAULT_LANGUAGE):
    builder = TEMPLATE_BUILDERS.get(template_key, build_from_moderno)
    builder(data, output_path, language_code=language_code)


_PDF_CHAR_REPLACEMENTS = {
    "‘": "'", "’": "'", "“": '"', "”": '"',
    "–": "-", "—": "-", "…": "...", "•": "-",
}


def _pdf_safe(text):
    """fpdf2's core (non-embedded) fonts only support Latin-1. Swap common
    "smart" punctuation for safe equivalents, and replace anything else
    outside Latin-1 rather than letting pdf.output() raise."""
    text = str(text or "")
    for src, dst in _PDF_CHAR_REPLACEMENTS.items():
        text = text.replace(src, dst)
    return text.encode("latin-1", errors="replace").decode("latin-1")


def build_linkedin_pdf(data, output_path, language_code=DEFAULT_LANGUAGE):
    if not isinstance(data, dict):
        data = {}
    from fpdf import FPDF

    labels = LANGUAGES.get(language_code, LANGUAGES[DEFAULT_LANGUAGE])["linkedin"]

    pdf = FPDF(format="Letter")
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.set_margins(22, 20, 22)
    pdf.add_page()

    def multi_cell_line(text, **kwargs):
        # multi_cell() doesn't reset the X cursor to the left margin after
        # rendering, so a later call can be left with ~0 width available
        # and raise "Not enough horizontal space to render a single
        # character". Reset X before every call.
        pdf.set_x(pdf.l_margin)
        pdf.multi_cell(0, text=text, **kwargs)

    pdf.set_font("Helvetica", "B", 18)
    multi_cell_line(_pdf_safe(labels["title"]), h=10)
    pdf.ln(6)

    def section(heading, body):
        pdf.set_font("Helvetica", "B", 13)
        multi_cell_line(_pdf_safe(heading), h=8)
        pdf.set_font("Helvetica", "", 11)
        multi_cell_line(_pdf_safe(body), h=6)
        pdf.ln(4)

    section(labels["headline"], data.get("headline") or "")
    section(labels["about"], data.get("about") or "")

    pdf.set_font("Helvetica", "B", 13)
    multi_cell_line(_pdf_safe(labels["recommendations"]), h=8)
    pdf.set_font("Helvetica", "", 11)
    recommendations = data.get("recommendations")
    if not isinstance(recommendations, list):
        recommendations = [recommendations] if recommendations else []
    for rec in recommendations:
        multi_cell_line(_pdf_safe(f"-  {rec}"), h=6)
        pdf.ln(1)

    pdf.output(output_path)
